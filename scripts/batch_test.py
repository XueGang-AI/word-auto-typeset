"""
End-to-end batch learning + typeset test.

Workflow:
1. Treat every .docx in `templates_dir` as a template example.
2. Analyze each (analyze_template) and aggregate the resulting configs.
3. For every .docx in `content_dir`, run the full pipeline:
   parse → AI recognize (with the aggregated style brief) → render.
4. Write the formatted outputs to `output_dir`.

Used to validate the batch-learning flow against a small set of
manually-edited ground-truth files in `test_files/`.
"""

from __future__ import annotations

import argparse
import json
import shutil
import sys
from pathlib import Path

# Allow running as `python scripts/batch_test.py` from project root.
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from app.config import TEMPLATES_DIR
from app.services.ai_recognizer import recognize_structure
from app.services.content_parser import parse_content
from app.services.renderer import render_document
from app.services.template_service import (
    aggregate_templates,
    analyze_template,
    build_style_brief,
    save_template,
)


def collect_docx(dir_path: Path) -> list[Path]:
    return sorted(p for p in dir_path.iterdir() if p.suffix.lower() == ".docx")


def step_learn(templates_dir: Path) -> tuple[list, object]:
    print(f"\n[1] Learning from templates in {templates_dir}")
    files = collect_docx(templates_dir)
    if not files:
        raise SystemExit(f"No .docx templates found in {templates_dir}")
    print(f"    found {len(files)} files: {[f.name for f in files]}")

    configs = []
    for f in files:
        doc_bytes = f.read_bytes()
        # save_template() persists under data/templates/<id>/; clear those
        # afterwards to keep the working tree clean.
        config = save_template(doc_bytes, f.name)
        configs.append(config)
        print(f"    - {f.name:50s} -> {config.template_id}  ({len(config.styles)} styles)")

    aggregated = aggregate_templates(configs)
    print(f"\n    aggregated config: {aggregated.template_id} "
          f"({len(aggregated.styles)} styles, "
          f"{len(aggregated.block_rules)} block rules)")

    # Persist the aggregated config so the UI / API can use it.
    tdir = TEMPLATES_DIR / aggregated.template_id
    tdir.mkdir(parents=True, exist_ok=True)
    (tdir / "template_config.json").write_text(
        aggregated.model_dump_json(indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    print(f"    saved to {tdir / 'template_config.json'}")
    print(f"\n    --- style brief ---")
    for line in build_style_brief(aggregated).splitlines():
        print(f"    {line}")

    return configs, aggregated


def step_process(
    content_dir: Path,
    output_dir: Path,
    aggregated,
) -> list[Path]:
    print(f"\n[2] Processing contents in {content_dir} -> {output_dir}")
    files = collect_docx(content_dir)
    if not files:
        raise SystemExit(f"No .docx content files found in {content_dir}")
    print(f"    found {len(files)} files: {[f.name for f in files]}")

    output_dir.mkdir(parents=True, exist_ok=True)

    style_brief = build_style_brief(aggregated)
    rendered: list[Path] = []
    for src in files:
        print(f"\n    -> {src.name}")
        paragraphs = parse_content(src)
        if not paragraphs:
            print("       skipped (no paragraphs)")
            continue
        print(f"       parsed {len(paragraphs)} paragraphs")

        content_doc = recognize_structure(
            paragraphs, src.name, style_brief=style_brief,
        )
        types = [p.para_type.value for article in content_doc.articles
                 for p in article.paragraphs]
        from collections import Counter
        print(f"       classified: {dict(Counter(types))}")
        if content_doc.warnings:
            for w in content_doc.warnings:
                print(f"       WARN: {w}")

        out_path = output_dir / f"{src.stem}_排版后.docx"
        render_document(content_doc, aggregated, out_path, None)
        print(f"       wrote {out_path}  ({out_path.stat().st_size:,} bytes)")
        rendered.append(out_path)

    return rendered


def step_compare(output_dir: Path, ground_truth_dir: Path) -> None:
    """Lightweight diff: list both sets and compare paragraph counts / types."""
    print(f"\n[3] Comparing outputs to ground truth in {ground_truth_dir}")
    if not ground_truth_dir.exists():
        print(f"    (skip — no ground-truth dir at {ground_truth_dir})")
        return
    outputs = {p.stem.replace("_排版后", ""): p for p in collect_docx(output_dir)}
    truths = collect_docx(ground_truth_dir)

    from docx import Document
    for truth in truths:
        key = truth.stem
        if key not in outputs:
            print(f"    [MISSING] output for {truth.name}")
            continue
        out = outputs[key]
        try:
            truth_doc = Document(str(truth))
            out_doc = Document(str(out))
        except Exception as e:
            print(f"    [ERROR] opening {truth.name}: {e}")
            continue
        n_truth = len(truth_doc.paragraphs)
        n_out = len(out_doc.paragraphs)
        n_img_truth = sum(1 for p in truth_doc.paragraphs if p._element.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
        ))
        n_img_out = sum(1 for p in out_doc.paragraphs if p._element.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
        ))
        flag = "OK " if n_out >= 0.5 * n_truth else "!! "
        print(f"    [{flag}] {key[:40]:40s}  truth={n_truth} paras / {n_img_truth} imgs  "
              f"out={n_out} paras / {n_img_out} imgs")


def main() -> None:
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("--templates", type=Path,
                   default=Path("/Users/xuegang/Desktop/My Project/test_files"))
    p.add_argument("--content", type=Path,
                   default=Path("/Users/xuegang/Desktop/My Project/input_files"))
    p.add_argument("--output", type=Path,
                   default=Path("/Users/xuegang/Desktop/My Project/output_files"))
    p.add_argument("--ground-truth", type=Path,
                   default=Path("/Users/xuegang/Desktop/My Project/test_files"),
                   help="Compare outputs to this directory (defaults to --templates).")
    p.add_argument("--cleanup", action="store_true",
                   help="Remove the per-file (tpl_*) template dirs after the run. "
                        "Aggregated templates are always kept.")
    args = p.parse_args()

    _, aggregated = step_learn(args.templates)

    rendered = step_process(args.content, args.output, aggregated)

    step_compare(args.output, args.ground_truth)

    if args.cleanup:
        # Remove the per-file (tpl_*) templates created during this run.
        # The aggregated (agg_*) template is kept so subsequent tools
        # (e.g. diff_truth.py, manual inspection) can reference it.
        from app.config import TEMPLATES_DIR
        removed = 0
        for child in TEMPLATES_DIR.iterdir():
            if child.is_dir() and child.name.startswith("tpl_"):
                shutil.rmtree(child, ignore_errors=True)
                removed += 1
        print(f"\n[cleanup] removed {removed} per-file template dirs from {TEMPLATES_DIR}"
              " (aggregated template preserved)")

    print(f"\nDone. {len(rendered)} files written to {args.output}")


if __name__ == "__main__":
    main()
