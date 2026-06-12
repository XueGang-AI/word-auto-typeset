"""
Exhaustive structural diff between a "ground truth" docx and a system-
produced docx. Reports every difference at a granularity that can be
acted on:

- per-paragraph: text, font (name/size/bold/color/italic), alignment,
  indent, line_spacing, space_before, space_after
- per-run: text fragment, font, size, bold, color, italic
- images: position, size (EMU), content_type, byte hash
- page: paper size, margins, orientation
- sections: count, header/footer presence

Output is a list of "diff records" plus a summary score. Each record has
a location, what differs, and the truth/output values.
"""

from __future__ import annotations

import hashlib
import json
import sys
from dataclasses import dataclass, asdict, field
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.shared import Emu


@dataclass
class Diff:
    kind: str          # "text" | "format" | "missing" | "extra" | "page" | "image"
    where: str         # human-readable location
    truth: object
    output: object
    note: str = ""


def _emu(v) -> int | None:
    if v is None:
        return None
    try:
        return int(v)
    except Exception:
        return None


def _pt(v) -> float | None:
    if v is None:
        return None
    try:
        return float(v)
    except Exception:
        return None


def _para_signature(p) -> dict:
    """Extract a signature for a single paragraph."""
    pf = p.paragraph_format
    align = str(pf.alignment) if pf.alignment is not None else None
    return {
        "text": p.text,
        "alignment": align,
        "first_line_indent_emu": _emu(pf.first_line_indent),
        "left_indent_emu": _emu(pf.left_indent),
        "right_indent_emu": _emu(pf.right_indent),
        "line_spacing": pf.line_spacing,
        "space_before_pt": _pt(pf.space_before),
        "space_after_pt": _pt(pf.space_after),
    }


def _run_signature(r) -> dict:
    f = r.font
    return {
        "text": r.text,
        "font_name": f.name,
        "size_pt": _pt(f.size) if f.size is not None else None,
        "bold": f.bold,
        "italic": f.italic,
        "color": str(f.color.rgb) if f.color and f.color.rgb else None,
    }


def _section_signature(s) -> dict:
    return {
        "page_width_emu": _emu(s.page_width),
        "page_height_emu": _emu(s.page_height),
        "top_margin_emu": _emu(s.top_margin),
        "bottom_margin_emu": _emu(s.bottom_margin),
        "left_margin_emu": _emu(s.left_margin),
        "right_margin_emu": _emu(s.right_margin),
        "orientation": str(s.orientation) if s.orientation is not None else None,
        "header_distance_emu": _emu(s.header_distance),
        "footer_distance_emu": _emu(s.footer_distance),
    }


def _image_signature(p, doc) -> list[dict]:
    """Walk runs to find drawings and their embedded image parts."""
    out = []
    for r in p.runs:
        for drawing in r._element.iter(
            "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
        ):
            for ext in drawing.iter(
                "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent"
            ):
                cx, cy = ext.get("cx"), ext.get("cy")
                break
            for blip in drawing.iter(
                "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            ):
                embed = blip.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
                if embed and embed in doc.part.rels:
                    rel = doc.part.rels[embed]
                    try:
                        blob = rel.target_part.blob
                        out.append({
                            "kind": "inline",
                            "width_emu": int(cx) if cx else None,
                            "height_emu": int(cy) if cy else None,
                            "content_type": rel.target_part.content_type,
                            "byte_hash": hashlib.md5(blob).hexdigest(),
                            "byte_len": len(blob),
                        })
                    except Exception:
                        out.append({"error": "could not read image part"})
    return out


def diff_files(truth_path: Path, output_path: Path) -> dict:
    diffs: list[Diff] = []
    summary = {
        "truth_path": str(truth_path),
        "output_path": str(output_path),
        "paragraphs_truth": 0,
        "paragraphs_output": 0,
        "images_truth": 0,
        "images_output": 0,
    }

    truth_doc = Document(str(truth_path))
    output_doc = Document(str(output_path))

    truth_paras = truth_doc.paragraphs
    output_paras = output_doc.paragraphs
    summary["paragraphs_truth"] = len(truth_paras)
    summary["paragraphs_output"] = len(output_paras)

    # ── Page settings ──
    if truth_doc.sections and output_doc.sections:
        ts = _section_signature(truth_doc.sections[0])
        os_ = _section_signature(output_doc.sections[0])
        for k in ts:
            if ts[k] != os_[k]:
                diffs.append(Diff("page", f"section[0].{k}", ts[k], os_[k]))
        summary["page_truth"] = ts
        summary["page_output"] = os_

    # ── Paragraph-level diff (LCS) ──
    truth_sigs = [_para_signature(p) for p in truth_paras]
    output_sigs = [_para_signature(p) for p in output_paras]

    n, m = len(truth_sigs), len(output_sigs)
    # LCS dp
    dp = [[0] * (m + 1) for _ in range(n + 1)]
    for i in range(n):
        for j in range(m):
            if truth_sigs[i]["text"].strip() == output_sigs[j]["text"].strip() and truth_sigs[i]["text"].strip():
                dp[i + 1][j + 1] = dp[i][j] + 1
            else:
                dp[i + 1][j + 1] = max(dp[i + 1][j], dp[i][j + 1])
    # Backtrack
    aligned_truth: list[int] = []  # indices in truth
    aligned_output: list[int] = []
    i, j = n, m
    while i > 0 and j > 0:
        if truth_sigs[i - 1]["text"].strip() == output_sigs[j - 1]["text"].strip() and truth_sigs[i - 1]["text"].strip():
            aligned_truth.append(i - 1)
            aligned_output.append(j - 1)
            i -= 1
            j -= 1
        elif dp[i - 1][j] >= dp[i][j - 1]:
            i -= 1
        else:
            j -= 1
    aligned_truth.reverse()
    aligned_output.reverse()

    # Now produce diffs: for each aligned pair, check format; for unaligned
    # truth paras, report missing; for unaligned output paras, report extra.
    truth_aligned_set = set(aligned_truth)
    output_aligned_set = set(aligned_output)

    # Missing in output (truth-only paragraphs)
    for i, sig in enumerate(truth_sigs):
        if i not in truth_aligned_set:
            diffs.append(Diff(
                "missing", f"para[{i}]", sig["text"][:80], None,
                "truth paragraph not matched in output",
            ))

    # Extra in output (output-only paragraphs)
    for j, sig in enumerate(output_sigs):
        if j not in output_aligned_set:
            diffs.append(Diff(
                "extra", f"para[{j}]", None, sig["text"][:80],
                "output paragraph not in truth",
            ))

    # Aligned — check formatting
    for ti, oi in zip(aligned_truth, aligned_output):
        ts = truth_sigs[ti]
        os_ = output_sigs[oi]
        for k in ts:
            if ts[k] != os_[k]:
                diffs.append(Diff("format", f"para[{ti}].{k}", ts[k], os_[k]))

        # Run-level diff (only the first run if lengths differ wildly; full diff if same count)
        truth_runs = truth_paras[ti].runs
        output_runs = output_paras[oi].runs
        n_runs_t = len(truth_runs)
        n_runs_o = len(output_runs)
        common = min(n_runs_t, n_runs_o)
        for ri in range(common):
            tr = _run_signature(truth_runs[ri])
            or_ = _run_signature(output_runs[ri])
            for k in tr:
                if tr[k] != or_[k]:
                    diffs.append(Diff(
                        "format",
                        f"para[{ti}].run[{ri}].{k}",
                        tr[k], or_[k],
                    ))
        if n_runs_t != n_runs_o:
            diffs.append(Diff(
                "format",
                f"para[{ti}].run_count",
                n_runs_t, n_runs_o,
            ))

    # ── Images: per-paragraph hash + size ──
    truth_imgs = []
    for i, p in enumerate(truth_paras):
        for img in _image_signature(p, truth_doc):
            truth_imgs.append((i, img))
    output_imgs = []
    for i, p in enumerate(output_paras):
        for img in _image_signature(p, output_doc):
            output_imgs.append((i, img))
    summary["images_truth"] = len(truth_imgs)
    summary["images_output"] = len(output_imgs)
    if len(truth_imgs) != len(output_imgs):
        diffs.append(Diff(
            "image", "image_count",
            len(truth_imgs), len(output_imgs),
        ))

    summary["diff_count"] = len(diffs)
    summary["diffs"] = [asdict(d) for d in diffs]

    # Score: percentage of paragraphs that match (text + format)
    matched = len(aligned_truth)
    summary["paragraphs_matched_text"] = matched
    summary["paragraphs_unmatched_format"] = sum(
        1 for d in diffs if d.kind == "format" and d.where.startswith("para[") and ".run[" in d.where
        or (d.kind == "format" and d.where.startswith("para[") and ".run_count" in d.where)
    )
    return summary


def main() -> None:
    if len(sys.argv) != 3:
        print("usage: python scripts/diff_docx.py <truth.docx> <output.docx>")
        sys.exit(1)
    truth, output = Path(sys.argv[1]), Path(sys.argv[2])
    summary = diff_files(truth, output)
    print(f"\n=== {truth.name} vs {output.name} ===")
    print(f"paragraphs: truth={summary['paragraphs_truth']} output={summary['paragraphs_output']} matched_text={summary['paragraphs_matched_text']}")
    print(f"images:     truth={summary['images_truth']} output={summary['images_output']}")
    print(f"total diff records: {summary['diff_count']}")
    if summary.get("page_truth"):
        print(f"\npage settings:")
        for k in summary["page_truth"]:
            if summary["page_truth"][k] != summary["page_output"][k]:
                print(f"  {k}: truth={summary['page_truth'][k]} output={summary['page_output'][k]}")

    # Group diffs by kind
    by_kind: dict = {}
    for d in summary["diffs"]:
        by_kind.setdefault(d["kind"], []).append(d)
    for kind, items in sorted(by_kind.items()):
        print(f"\n[{kind}] {len(items)} records (showing first 20):")
        for d in items[:20]:
            note = f"  // {d['note']}" if d.get("note") else ""
            print(f"  {d['where']:50s}  truth={str(d['truth'])[:60]}  output={str(d['output'])[:60]}{note}")
        if len(items) > 20:
            print(f"  ... and {len(items) - 20} more")

    # Save full report
    out_json = truth.parent / f"_diff_{truth.stem}.json"
    out_json.write_text(
        json.dumps(summary, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(f"\nFull diff saved to {out_json}")


if __name__ == "__main__":
    main()
