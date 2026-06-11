"""
Compare the two processed outputs to the corresponding ground-truth files
in test_files/. Produces a per-file report:

- structural (paragraph / image / table counts)
- per-paragraph text diff
- AI classification diff (run recognition on both truth and output, then
  compare the resulting types)

The point is to expose where the batch-learning pipeline diverges from
the manually-edited ground truth.
"""

from __future__ import annotations

import json
import sys
from collections import Counter
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document

from app.services.ai_recognizer import recognize_structure
from app.services.content_parser import parse_content
from app.services.template_service import build_style_brief, get_template, list_templates


TRUTH_DIR = Path("/Users/xuegang/Desktop/My Project/test_files")
INPUT_DIR = Path("/Users/xuegang/Desktop/My Project/input_files")
OUTPUT_DIR = Path("/Users/xuegang/Desktop/My Project/output_files")
REPORT_PATH = Path("/Users/xuegang/Desktop/My Project/output_files/_REPORT.md")

# Find the most recent aggregated template.
agg_cfgs = [t for t in list_templates() if t.template_id.startswith("agg_")]
if not agg_cfgs:
    print("No aggregated template found — run batch_test.py first.")
    sys.exit(0)
latest_agg = sorted(agg_cfgs, key=lambda t: t.template_id, reverse=True)[0]
print(f"Using aggregated template: {latest_agg.template_id}")
agg_config = get_template(latest_agg.template_id)
style_brief = build_style_brief(agg_config)


def structural_stats(path: Path) -> dict:
    doc = Document(str(path))
    paras = doc.paragraphs
    images = sum(
        1 for p in paras
        if p._element.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
        )
    )
    tables = len(doc.tables)
    return {
        "paragraphs": len(paras),
        "images": images,
        "tables": tables,
    }


def classify_file(path: Path, label: str) -> list[str]:
    paragraphs = parse_content(path)
    if not paragraphs:
        return []
    content_doc = recognize_structure(paragraphs, path.name, style_brief=style_brief)
    types = [
        p.para_type.value
        for article in content_doc.articles
        for p in article.paragraphs
    ]
    print(f"  {label}: classified {len(types)} paragraphs "
          f"({dict(Counter(types))})")
    return types


# Find input files and their matching truth / output by stem.
inputs = sorted(INPUT_DIR.glob("*.docx"))
lines = [
    "# 批量学习测试 — 报告",
    "",
    f"- 聚合模板: `{latest_agg.template_id}`",
    f"- 模板数: {len([t for t in list_templates() if t.template_id.startswith('tpl_')])}",
    f"- 输入/输出目录: `{INPUT_DIR}` → `{OUTPUT_DIR}`",
    f"- 真值目录: `{TRUTH_DIR}`",
    "",
    "## 结构对比",
    "",
    "| 文件 | 真值段落 | 真值图片 | 真值表格 | 输出段落 | 输出图片 | 输出表格 |",
    "|---|---|---|---|---|---|---|",
]

for inp in inputs:
    truth = TRUTH_DIR / inp.name
    output = OUTPUT_DIR / f"{inp.stem}_排版后.docx"
    if not truth.exists():
        continue
    if not output.exists():
        lines.append(f"| {inp.name} | — | — | — | MISSING | — | — |")
        continue
    ts = structural_stats(truth)
    os_ = structural_stats(output)
    lines.append(
        f"| {inp.name[:50]} | {ts['paragraphs']} | {ts['images']} | {ts['tables']} "
        f"| {os_['paragraphs']} | {os_['images']} | {os_['tables']} |"
    )

lines += ["", "## AI 分类对比（同样 prompt 跑真值 vs 输出）", ""]
for inp in inputs:
    truth = TRUTH_DIR / inp.name
    output = OUTPUT_DIR / f"{inp.stem}_排版后.docx"
    if not truth.exists() or not output.exists():
        continue
    lines.append(f"### {inp.name}")
    lines.append("")
    truth_types = classify_file(truth, "truth")
    out_types = classify_file(output, "output")
    lines.append(f"- 真值类型分布: `{dict(Counter(truth_types))}`")
    lines.append(f"- 输出类型分布: `{dict(Counter(out_types))}`")
    if truth_types and out_types:
        agreement = sum(1 for a, b in zip(truth_types, out_types) if a == b)
        lines.append(f"- 同位置类型一致: {agreement}/{min(len(truth_types), len(out_types))}")
    lines.append("")

REPORT_PATH.write_text("\n".join(lines), encoding="utf-8")
print(f"\nReport written to {REPORT_PATH}")
