#!/usr/bin/env python3
"""
格式对齐工具：将内容文档的格式对齐到模板文档。
- 提取模板文档的页面设置和段落格式特征
- 智能检测内容文档的段落类型
- 应用模板格式，输出对齐后的 Word 文档
"""

import argparse
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ── 工具函数 ───────────────────────────────────────────

def is_empty_paragraph(para) -> bool:
    """判断段落是否为空（无文本且无图片）。"""
    if para.text.strip():
        return False
    # 检查是否包含图片
    for run in para.runs:
        drawings = run._element.findall(qn('w:drawing'))
        if drawings:
            return False
        inline_shapes = run._element.findall(qn('wp:inline'))
        if inline_shapes:
            return False
    # Also check for direct drawings
    drawings = para._element.findall(qn('w:drawing'))
    if drawings:
        return False
    return True


def has_image(para) -> bool:
    """判断段落是否包含图片。"""
    for run in para.runs:
        if run._element.findall(qn('w:drawing')) or run._element.findall(qn('wp:inline')):
            return True
    if para._element.findall(qn('w:drawing')):
        return True
    return False


def is_blue_color(rgb) -> bool:
    """判断颜色是否为蓝色系。"""
    if rgb is None:
        return False
    if isinstance(rgb, str):
        rgb_str = rgb.lstrip('#')
        r, g, b = int(rgb_str[0:2], 16), int(rgb_str[2:4], 16), int(rgb_str[4:6], 16)
    else:
        r, g, b = int(rgb[0:2], 16), int(rgb[2:4], 16), int(rgb[4:6], 16)
    return b > r + 40 and b > g + 20


def is_red_color(rgb) -> bool:
    """判断颜色是否为红色系。"""
    if rgb is None:
        return False
    if isinstance(rgb, str):
        rgb_str = rgb.lstrip('#')
        r, g, b = int(rgb_str[0:2], 16), int(rgb_str[2:4], 16), int(rgb_str[4:6], 16)
    else:
        r, g, b = int(rgb[0:2], 16), int(rgb[2:4], 16), int(rgb[4:6], 16)
    return r > g + 40 and r > b + 40


def get_first_run_color(para):
    """获取段落第一个有颜色的 run 的颜色。"""
    for run in para.runs:
        if run.font.color and run.font.color.rgb:
            return str(run.font.color.rgb)
    return None


def get_first_run_bold(para):
    """获取段落第一个 run 是否加粗。"""
    for run in para.runs:
        if run.font.bold is not None:
            return run.font.bold
    return None


def looks_like_person_name(text: str) -> bool:
    """判断文本是否像人名（2-4个汉字，无标点）。"""
    text = text.strip()
    if not (2 <= len(text) <= 4):
        return False
    # 全是中文字符
    if not all('一' <= c <= '鿿' for c in text):
        return False
    return True


def is_affiliation(text: str) -> bool:
    """判断是否为作者单位/执笔人信息。"""
    text = text.strip()
    return text.startswith('（') and ('作者为' in text or '执笔人' in text or '作者单位' in text)


# ── 模板格式提取 ────────────────────────────────────────

def extract_template_profiles(doc: Document) -> dict:
    """
    从模板文档动态提取格式配置。
    支持替换模板：不同模板自动提取不同的格式参数。
    """
    profiles = {}

    # ── 页面设置（取第一节） ──
    section = doc.sections[0]
    profiles['page_width'] = section.page_width
    profiles['page_height'] = section.page_height
    profiles['top_margin'] = section.top_margin
    profiles['bottom_margin'] = section.bottom_margin
    profiles['left_margin'] = section.left_margin
    profiles['right_margin'] = section.right_margin

    # ── 从模板段落中动态采样格式 ──
    main_title_fmt = None      # 第1段 → 主标题
    body_text_fmt = None       # 有首行缩进(2字符左右) → 正文
    affiliation_fmt = None     # 有大缩进 → 作者单位
    default_run_fmt = None     # 兜底：正文段落的首个 run 格式

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        pf = para.paragraph_format
        first_run = para.runs[0] if para.runs else None
        if first_run is None:
            continue

        # 从该 run 提取基本信息
        run_info = {
            'font_name': first_run.font.name,
            'font_size': first_run.font.size,
            'font_bold': first_run.font.bold,
            'first_line_indent': pf.first_line_indent,
            'alignment': pf.alignment,
        }

        # 第 1 个非空段落 → 主标题
        if main_title_fmt is None:
            main_title_fmt = dict(run_info)

        # 正文：首行缩进在 2-4 个中文字符范围内（14pt 字体下约 355600-711200 EMU）
        if pf.first_line_indent and 150000 < pf.first_line_indent < 800000:
            if body_text_fmt is None:
                body_text_fmt = dict(run_info)
            # 同时作为兜底格式
            if default_run_fmt is None:
                default_run_fmt = dict(run_info)

        # 作者单位：大缩进
        if pf.first_line_indent and pf.first_line_indent >= 800000:
            if affiliation_fmt is None:
                affiliation_fmt = dict(run_info)

        # 兜底：取任意正文段落的格式
        if default_run_fmt is None and pf.first_line_indent is None:
            default_run_fmt = dict(run_info)

    # ── 如果模板没有缩进段落（如纯标题型模板），用任意段落兜底 ──
    if default_run_fmt is None:
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text or not para.runs:
                continue
            r = para.runs[0]
            default_run_fmt = {
                'font_name': r.font.name,
                'font_size': r.font.size,
                'font_bold': r.font.bold,
                'first_line_indent': para.paragraph_format.first_line_indent,
                'alignment': para.paragraph_format.alignment,
            }
            break

    # ── 最终兜底 ──
    if default_run_fmt is None:
        default_run_fmt = {
            'font_name': '仿宋', 'font_size': Pt(14), 'font_bold': False,
            'first_line_indent': None, 'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        }

    # ── 组装 profiles ──
    # body_text：用模板正文格式（含缩进）
    body_final = dict(body_text_fmt or default_run_fmt)
    body_final['line_spacing'] = 1.5
    body_final.setdefault('alignment', WD_ALIGN_PARAGRAPH.JUSTIFY)
    profiles['body_text'] = body_final

    # main_title：用模板第 1 段格式
    mt = dict(main_title_fmt or default_run_fmt)
    mt['line_spacing'] = 1.5
    mt['alignment'] = mt.get('alignment') or WD_ALIGN_PARAGRAPH.LEFT
    mt.pop('first_line_indent', None)
    profiles['main_title'] = mt

    # affiliation：用模板作者单位格式（含大缩进）
    aff = dict(affiliation_fmt or default_run_fmt)
    aff['line_spacing'] = 1.5
    aff['alignment'] = aff.get('alignment') or WD_ALIGN_PARAGRAPH.LEFT
    profiles['affiliation'] = aff

    # 其余类型：基础格式（仿宋 14pt，无缩进，不加粗）
    base = {
        'font_name': default_run_fmt['font_name'],
        'font_size': default_run_fmt['font_size'],
        'line_spacing': 1.5,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
    }
    # 如果模板默认字体为 None，用仿宋兜底
    if base['font_name'] is None:
        base['font_name'] = '仿宋'
    if base['font_size'] is None:
        base['font_size'] = Pt(14)

    profiles['article_title'] = dict(base)
    profiles['section_header'] = dict(base)
    profiles['section_header']['font_bold'] = True
    profiles['author_name'] = dict(base)
    profiles['tag_label'] = dict(base)

    return profiles


# ── 段落类型检测 ────────────────────────────────────────

def detect_paragraph_type(para, is_first_content: bool, has_seen_body: bool) -> str:
    """
    智能检测内容文档中段落的类型。
    has_seen_body: 是否已出现过正文段落（用于区分前置元数据和正文标题）。
    返回类型标识字符串。
    """
    text = para.text.strip()

    if not text:
        if has_image(para):
            return 'image'
        return 'empty'

    # 第1个非空段落 → 主标题
    if is_first_content:
        return 'main_title'

    # 作者单位 / 执笔人
    if is_affiliation(text):
        return 'affiliation'

    # 长段落 → 正文（优先判定，避免被颜色规则误判）
    if len(text) > 50:
        return 'body_text'

    # 短段落：按颜色进一步细分
    color = get_first_run_color(para)

    if color:
        if is_blue_color(color):
            # 正文区域内的蓝色长标题才是层级小标题
            # 正文之前的蓝色短段落是作者信息等元数据
            if len(text) > 15 and has_seen_body:
                return 'section_header'
            return 'tag_label'
        if is_red_color(color):
            if len(text) > 8:
                return 'article_title'
            return 'tag_label'

    # 短文本，无特殊颜色
    if looks_like_person_name(text):
        return 'author_name'

    # 其余短文本 → 标签/导读
    return 'tag_label'


# ── 格式应用 ────────────────────────────────────────────

def apply_format(para, para_type: str, profiles: dict):
    """
    将模板格式应用到指定段落。
    对齐：字体、字号、行距、缩进、对齐方式。
    保留原有颜色和粗体标记。
    """
    profile = profiles.get(para_type, profiles['body_text'])

    # === 段落级别格式 ===
    pf = para.paragraph_format

    # 行距 1.5 倍
    pf.line_spacing = 1.5

    # 对齐方式
    if 'alignment' in profile and profile['alignment'] is not None:
        pf.alignment = profile['alignment']

    # 首行缩进
    if 'first_line_indent' in profile:
        pf.first_line_indent = profile['first_line_indent']
    else:
        pf.first_line_indent = None

    # 清除段前段后间距（对齐紧凑排版）
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # === Run 级别格式 ===
    for run in para.runs:
        # 字体名称
        if 'font_name' in profile and profile['font_name']:
            run.font.name = profile['font_name']
            # 设置中文字体
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                from lxml import etree
                rFonts = etree.SubElement(rPr, qn('w:rFonts'))
            rFonts.set(qn('w:eastAsia'), profile['font_name'])

        # 字号
        if 'font_size' in profile and profile['font_size']:
            run.font.size = profile['font_size']

        # 粗体控制：
        # - section_header: 强制加粗
        # - body_text: 保留原文档的粗体（可能是内容强调）
        # - 其余类型: 取消加粗，对齐模板
        if para_type == 'section_header':
            run.font.bold = True
        elif para_type != 'body_text':
            run.font.bold = False


def apply_page_settings(doc: Document, profiles: dict):
    """应用模板的页面设置。"""
    for section in doc.sections:
        section.page_width = profiles['page_width']
        section.page_height = profiles['page_height']
        section.top_margin = profiles['top_margin']
        section.bottom_margin = profiles['bottom_margin']
        section.left_margin = profiles['left_margin']
        section.right_margin = profiles['right_margin']


# ── 主流程 ──────────────────────────────────────────────

def format_document(template_path: Path, content_path: Path, output_path: Path) -> dict:
    """
    将内容文档格式对齐到模板文档。
    返回统计信息。
    """
    template_doc = Document(str(template_path))
    content_doc = Document(str(content_path))

    # 1. 提取模板格式
    profiles = extract_template_profiles(template_doc)

    # 2. 应用页面设置
    apply_page_settings(content_doc, profiles)

    # 3. 遍历内容段落，检测类型并应用格式
    stats = {
        'main_title': 0,
        'article_title': 0,
        'section_header': 0,
        'author_name': 0,
        'tag_label': 0,
        'body_text': 0,
        'affiliation': 0,
        'image': 0,
        'empty_removed': 0,
        'total_original': len(content_doc.paragraphs),
    }

    found_first_content = False
    has_seen_body = False
    paragraphs_to_remove = []

    for i, para in enumerate(content_doc.paragraphs):
        if is_empty_paragraph(para) and not has_image(para):
            paragraphs_to_remove.append(para)
            stats['empty_removed'] += 1
            continue

        para_type = detect_paragraph_type(para, not found_first_content, has_seen_body)

        if para_type == 'empty':
            paragraphs_to_remove.append(para)
            stats['empty_removed'] += 1
            continue

        if not found_first_content and para_type not in ('empty', 'image'):
            found_first_content = True

        if para_type == 'body_text':
            has_seen_body = True

        stats[para_type] = stats.get(para_type, 0) + 1
        apply_format(para, para_type, profiles)

    # 4. 删除空段落
    for para in paragraphs_to_remove:
        try:
            p_element = para._element
            p_element.getparent().remove(p_element)
        except Exception:
            pass

    # 5. 保存
    output_path.parent.mkdir(parents=True, exist_ok=True)
    content_doc.save(str(output_path))

    return stats


def main() -> int:
    parser = argparse.ArgumentParser(
        description="格式对齐工具：将内容文档的格式对齐到模板文档"
    )
    parser.add_argument(
        "--template", required=True,
        help="模板文档路径（提供目标格式）"
    )
    parser.add_argument(
        "--content", required=True,
        help="内容文档路径（需要格式化的文档）"
    )
    parser.add_argument(
        "--output", default=None,
        help="输出路径（默认：内容文档名_格式对齐.docx）"
    )
    args = parser.parse_args()

    template_path = Path(args.template).expanduser().resolve()
    content_path = Path(args.content).expanduser().resolve()

    if not template_path.exists():
        print(f"[ERROR] 模板文件不存在: {template_path}")
        return 2
    if not content_path.exists():
        print(f"[ERROR] 内容文件不存在: {content_path}")
        return 2

    if args.output:
        output_path = Path(args.output).expanduser().resolve()
    else:
        output_path = content_path.parent / f"{content_path.stem}_格式对齐.docx"

    print(f"[INFO] 模板: {template_path.name}")
    print(f"[INFO] 内容: {content_path.name}")
    print(f"[INFO] 输出: {output_path.name}")
    print()

    try:
        stats = format_document(template_path, content_path, output_path)
    except Exception as e:
        print(f"[ERROR] 转换失败: {e}")
        import traceback
        traceback.print_exc()
        return 1

    print("[OK] 格式对齐完成！")
    print(f"     原段落数: {stats['total_original']}")
    print(f"     删除空行: {stats['empty_removed']}")
    print(f"     主标题:   {stats.get('main_title', 0)}")
    print(f"     章节标题: {stats.get('article_title', 0)}")
    print(f"     层级标题: {stats.get('section_header', 0)}")
    print(f"     作者名:   {stats.get('author_name', 0)}")
    print(f"     标签:     {stats.get('tag_label', 0)}")
    print(f"     正文:     {stats.get('body_text', 0)}")
    print(f"     作者单位: {stats.get('affiliation', 0)}")
    print(f"     图片:     {stats.get('image', 0)}")
    print(f"     输出:     {output_path}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
