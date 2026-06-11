"""
Word renderer: applies a template configuration to structured content
and generates a formatted Word document.
"""

from __future__ import annotations

import copy
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor

from app.schemas.models import (
    BlockRule,
    ContentDocument,
    ContentParagraph,
    ParagraphType,
    TemplateConfig,
)
from app.utils.word_utils import has_image


# Alignment mapping
_ALIGN_MAP = {
    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "0": WD_ALIGN_PARAGRAPH.LEFT,
    "1": WD_ALIGN_PARAGRAPH.CENTER,
    "2": WD_ALIGN_PARAGRAPH.RIGHT,
    "3": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def render_document(
    content: ContentDocument,
    template: TemplateConfig,
    output_path: Path,
    template_docx_path: Path | None = None,
) -> Path:
    """
    Render content into a formatted Word document using a template.

    Strategy:
    1. Start from a copy of the template docx (to preserve headers/footers/styles)
       OR create a blank document with template page settings.
    2. Clear existing content.
    3. For each article and each paragraph, apply the matching template style.
    4. Handle page breaks between articles.
    5. Handle images, tables.
    """
    # ── Create base document ──
    if template_docx_path and template_docx_path.exists():
        doc = Document(str(template_docx_path))
        # Remove existing paragraphs (keep section properties)
        _clear_document_body(doc)
    else:
        doc = Document()

    # ── Apply page settings ──
    _apply_page_settings(doc, template)

    # ── Render each article ──
    total_articles = len(content.articles)
    for art_idx, article in enumerate(content.articles):
        if art_idx > 0:
            # Page break between articles
            doc.add_page_break()

        for cp in article.paragraphs:
            # Check block rule
            rule = template.block_rules.get(cp.para_type.value, BlockRule.optional)
            if rule == BlockRule.skip:
                continue

            # Render based on type
            if cp.is_table and cp.table_data:
                _render_table(doc, cp, template)
            elif cp.has_image and cp.para_type == ParagraphType.image:
                _render_image_placeholder(doc, cp, template)
            else:
                _render_paragraph(doc, cp, template)

    # ── Save ──
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ── Paragraph Rendering ───────────────────────────────────

def _render_paragraph(doc: Document, cp: ContentParagraph, template: TemplateConfig) -> None:
    """Render a text paragraph with template styling."""
    para_type_value = cp.para_type.value
    style = template.styles.get(para_type_value)
    if not style:
        # Fallback to body_text style
        style = template.styles.get("body_text")
    if not style:
        style = template.styles.get("main_title")

    # Ensure font_name is never None — fall back to body_text font
    if style and style.font_name is None:
        body_style = template.styles.get("body_text")
        if body_style and body_style.font_name:
            style = style.model_copy(update={"font_name": body_style.font_name})

    para = doc.add_paragraph()
    pf = para.paragraph_format

    # ── Alignment ──
    if style.alignment:
        pf.alignment = _ALIGN_MAP.get(style.alignment, WD_ALIGN_PARAGRAPH.LEFT)

    # ── Line spacing ──
    if style.line_spacing:
        pf.line_spacing = style.line_spacing

    # ── Space before/after ──
    if style.space_before_pt is not None:
        pf.space_before = Pt(style.space_before_pt)
    if style.space_after_pt is not None:
        pf.space_after = Pt(style.space_after_pt)

    # ── First line indent ──
    if style.first_line_indent_emu is not None:
        pf.first_line_indent = Emu(style.first_line_indent_emu)
    elif cp.para_type == ParagraphType.body_text:
        # Body text should have indent unless template says otherwise
        pass

    # ── Add runs ──
    if cp.runs:
        # Preserve original run structure
        for run_data in cp.runs:
            run = para.add_run(run_data.get("text", ""))
            _apply_run_format(run, style, cp.para_type, run_data)
    else:
        # Simple single-run paragraph
        run = para.add_run(cp.text)
        _apply_run_format(run, style, cp.para_type)


def _apply_run_format(run, style, para_type: ParagraphType, run_data: dict | None = None) -> None:
    """Apply font formatting to a run."""
    # Font name
    if style.font_name:
        run.font.name = style.font_name
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            from lxml import etree
            rFonts = etree.SubElement(rPr, qn("w:rFonts"))
        rFonts.set(qn("w:eastAsia"), style.font_name)

    # Font size
    if style.font_size_emu:
        run.font.size = Emu(style.font_size_emu)
    elif style.font_size_pt:
        run.font.size = Pt(style.font_size_pt)

    # Bold
    if style.font_bold is not None:
        # For body_text, preserve original bold (content emphasis)
        if para_type == ParagraphType.body_text and run_data:
            original_bold = run_data.get("bold")
            if original_bold is not None:
                run.font.bold = original_bold
            else:
                run.font.bold = style.font_bold
        elif para_type == ParagraphType.section_header:
            run.font.bold = True
        else:
            run.font.bold = style.font_bold

    # Italic
    if style.font_italic is not None:
        run.font.italic = style.font_italic

    # Color — apply style color, but preserve original colors for body text
    if run_data and run_data.get("color") and para_type == ParagraphType.body_text:
        try:
            run.font.color.rgb = RGBColor.from_string(run_data["color"])
        except Exception:
            pass
    elif style.font_color:
        try:
            run.font.color.rgb = RGBColor.from_string(style.font_color)
        except Exception:
            pass


# ── Table Rendering ───────────────────────────────────────

def _render_table(doc: Document, cp: ContentParagraph, template: TemplateConfig) -> None:
    """Render a table."""
    if not cp.table_data:
        return
    rows = len(cp.table_data)
    cols = max((len(row) for row in cp.table_data), default=0)
    if rows == 0 or cols == 0:
        return

    table = doc.add_table(rows=rows, cols=cols, style="Table Grid")

    body_style = template.styles.get("body_text")
    font_name = body_style.font_name if body_style else "仿宋"
    font_size = Pt(body_style.font_size_pt) if body_style and body_style.font_size_pt else Pt(11)

    for i, row_data in enumerate(cp.table_data):
        for j, cell_text in enumerate(row_data):
            if j < cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    for run in para.runs:
                        run.font.name = font_name
                        run.font.size = font_size
                        rPr = run._element.get_or_add_rPr()
                        rFonts = rPr.find(qn("w:rFonts"))
                        if rFonts is None:
                            from lxml import etree
                            rFonts = etree.SubElement(rPr, qn("w:rFonts"))
                        rFonts.set(qn("w:eastAsia"), font_name)

    # Add an empty paragraph after the table for spacing
    doc.add_paragraph()


# ── Image Rendering ───────────────────────────────────────

def _apply_image_paragraph_format(para, img_style) -> None:
    """Apply spacing and alignment from image style to a paragraph."""
    pf = para.paragraph_format
    # Alignment — default to CENTER for images
    if img_style and img_style.alignment:
        pf.alignment = _ALIGN_MAP.get(img_style.alignment, WD_ALIGN_PARAGRAPH.CENTER)
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Line spacing
    if img_style and img_style.line_spacing:
        pf.line_spacing = img_style.line_spacing
    # Space before/after
    if img_style:
        if img_style.space_before_pt is not None:
            pf.space_before = Pt(img_style.space_before_pt)
        if img_style.space_after_pt is not None:
            pf.space_after = Pt(img_style.space_after_pt)
    # Never apply first-line indent to images


def _render_image_placeholder(doc: Document, cp: ContentParagraph, template: TemplateConfig) -> None:
    """
    Render images from the content paragraph into the document.

    Extracts image bytes from cp.images and inserts them using python-docx's
    add_picture, which handles relationship plumbing automatically.
    Falls back to a [图片] text placeholder if image extraction fails.
    """
    import io

    img_style = template.styles.get("image", template.styles.get("body_text"))
    images = cp.images or []

    if images:
        # Insert each image
        for img_info in images:
            img_bytes = img_info.get("bytes")
            if not img_bytes:
                continue

            para = doc.add_paragraph()
            _apply_image_paragraph_format(para, img_style)

            try:
                width_emu = img_info.get("width_emu")
                height_emu = img_info.get("height_emu")

                run = para.add_run()
                if width_emu and height_emu:
                    run.add_picture(
                        io.BytesIO(img_bytes),
                        width=Emu(width_emu),
                        height=Emu(height_emu),
                    )
                else:
                    run.add_picture(io.BytesIO(img_bytes))
            except Exception:
                # Fallback: show placeholder
                run = para.add_run("[图片]")
                if img_style:
                    if img_style.font_name:
                        run.font.name = img_style.font_name
                    if img_style.font_size_pt:
                        run.font.size = Pt(img_style.font_size_pt)
    else:
        # No image data extracted — show text placeholder
        para = doc.add_paragraph()
        _apply_image_paragraph_format(para, img_style)

        if cp.runs and any(r.get("text", "").strip() for r in cp.runs):
            # Has meaningful text runs — render them
            pass  # already added paragraph, runs handled elsewhere
        else:
            para.add_run("[图片]")


# ── Helpers ───────────────────────────────────────────────

def _clear_document_body(doc: Document) -> None:
    """Remove all paragraphs from the document body, preserving sections."""
    body = doc.element.body
    elements_to_remove = []
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            elements_to_remove.append(child)
        elif tag == "tbl":
            elements_to_remove.append(child)

    for elem in elements_to_remove:
        body.remove(elem)


def _apply_page_settings(doc: Document, template: TemplateConfig) -> None:
    """Apply template page settings to all sections."""
    ps = template.page_settings
    if not ps:
        return

    for section in doc.sections:
        if ps.width_emu:
            section.page_width = Emu(ps.width_emu)
        if ps.height_emu:
            section.page_height = Emu(ps.height_emu)
        if ps.top_margin_emu is not None:
            section.top_margin = Emu(ps.top_margin_emu)
        if ps.bottom_margin_emu is not None:
            section.bottom_margin = Emu(ps.bottom_margin_emu)
        if ps.left_margin_emu is not None:
            section.left_margin = Emu(ps.left_margin_emu)
        if ps.right_margin_emu is not None:
            section.right_margin = Emu(ps.right_margin_emu)
