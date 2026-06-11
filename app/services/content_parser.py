"""
Content parser: extracts structured content from a Word document.
Parses paragraphs, tables, images, and cleans up noise.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from app.schemas.models import ContentParagraph, ParagraphType
from app.utils.word_utils import (
    count_images,
    get_run_font_info,
    has_image,
    is_empty_paragraph,
    looks_like_ad,
)


def parse_content(file_path: Path) -> list[dict[str, Any]]:
    """
    Parse a Word document into a list of paragraph dicts.

    Each paragraph dict includes:
    - index: original position
    - text: the full paragraph text
    - has_image: bool
    - image_count: number of images
    - images: list of dicts with image bytes, content_type, width_emu, height_emu
    - is_table: bool
    - formatting hints from the first run
    - original runs data for preservation
    """
    doc = Document(str(file_path))
    paragraphs: list[dict[str, Any]] = []
    para_index = 0

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Find the corresponding python-docx paragraph
            if para_index >= len(doc.paragraphs):
                continue
            para = doc.paragraphs[para_index]
            para_index += 1

            text = para.text.strip()

            # Skip empty paragraphs (pure whitespace, no images)
            if is_empty_paragraph(para) and not has_image(para):
                continue

            # Skip ads/spam
            if text and looks_like_ad(text):
                continue

            # Collect run information and extract images
            runs_data = []
            extracted_images: list[dict[str, Any]] = []
            has_hyperlink = False
            for run in para.runs:
                run_info = get_run_font_info(run)
                run_info["text"] = run.text
                # Check for hyperlink
                parent = run._element.getparent()
                while parent is not None:
                    if parent.tag.endswith("}hyperlink"):
                        has_hyperlink = True
                        break
                    parent = parent.getparent()
                run_info["is_hyperlink"] = has_hyperlink
                runs_data.append(run_info)

                # Extract images from this run
                run_images = _extract_images_from_run(run, doc)
                extracted_images.extend(run_images)

            # Also extract images directly from the paragraph element
            if not extracted_images:
                para_images = _extract_images_from_element(para._element, doc)
                extracted_images.extend(para_images)

            # Skip meaningless hyperlink-only empty paragraphs
            if not text and has_hyperlink and not extracted_images:
                continue

            paragraphs.append({
                "index": len(paragraphs),
                "text": text,
                "has_image": has_image(para) or bool(extracted_images),
                "image_count": count_images(para) or len(extracted_images),
                "images": extracted_images,
                "is_table": False,
                "runs": runs_data,
                "format_hints": {
                    "first_run": runs_data[0] if runs_data else {},
                    "is_centered": _is_centered(para),
                    "is_bold": runs_data[0].get("bold", False) if runs_data else False,
                },
            })

        elif tag == "tbl":
            # Table handling
            table_data = _parse_table_element(element, doc)
            paragraphs.append({
                "index": len(paragraphs),
                "text": table_data.get("text_summary", ""),
                "has_image": False,
                "image_count": 0,
                "images": [],
                "is_table": True,
                "table_data": table_data.get("rows", []),
                "runs": [],
                "format_hints": {"is_table": True},
            })

    return paragraphs


def _extract_images_from_run(run, doc: Document) -> list[dict[str, Any]]:
    """Extract image data from a run's drawing elements."""
    images = []
    nsmap = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    }

    for drawing_elem in run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"):
        images.extend(_extract_images_from_drawing(drawing_elem, doc, nsmap))

    return images


def _extract_images_from_element(elem, doc: Document) -> list[dict[str, Any]]:
    """Extract image data directly from paragraph element's drawing children."""
    images = []
    nsmap = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    }

    for drawing_elem in elem.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"):
        images.extend(_extract_images_from_drawing(drawing_elem, doc, nsmap))

    return images


def _extract_images_from_drawing(drawing_elem, doc: Document, nsmap: dict) -> list[dict[str, Any]]:
    """Extract image data from a w:drawing element."""
    images = []

    for container_tag in ("{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline",
                          "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor"):
        for container in drawing_elem.findall(container_tag):
            # Get extent (dimensions)
            extent = container.find("{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent")
            width_emu = int(extent.get("cx")) if extent is not None else None
            height_emu = int(extent.get("cy")) if extent is not None else None

            # Find blip elements (image references)
            for blip in container.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"):
                embed_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                if embed_id and embed_id in doc.part.rels:
                    rel = doc.part.rels[embed_id]
                    try:
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        content_type = image_part.content_type
                        images.append({
                            "bytes": image_bytes,
                            "content_type": content_type,
                            "width_emu": width_emu,
                            "height_emu": height_emu,
                        })
                    except Exception:
                        pass

    return images


def _is_centered(para) -> bool:
    """Check if paragraph has center alignment."""
    if para.paragraph_format.alignment is not None:
        return "CENTER" in str(para.paragraph_format.alignment).upper()
    return False


def _parse_table_element(tbl_element, doc: Document) -> dict[str, Any]:
    """Parse a w:tbl element into structured data."""
    rows = []
    text_parts = []
    for row_elem in tbl_element.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr"):
        cells = []
        for cell_elem in row_elem.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc"):
            cell_text_parts = []
            for p_elem in cell_elem.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"):
                t_texts = []
                for t_elem in p_elem.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                    if t_elem.text:
                        t_texts.append(t_elem.text)
                cell_text_parts.append("".join(t_texts))
            cell_text = " ".join(cell_text_parts).strip()
            cells.append(cell_text)
            text_parts.append(cell_text)
        rows.append(cells)
    return {"rows": rows, "text_summary": " | ".join(filter(None, text_parts))}
